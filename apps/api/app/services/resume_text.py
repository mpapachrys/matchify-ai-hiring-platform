"""Package an uploaded resume for the parser.

Downloading lives on the platform side, not behind the AI seam: the parser is
handed bytes or text and never needs object-storage credentials.

PDFs are deliberately *not* reduced to text here — see `build_source`. Other
formats have no layout to lose, so those are extracted as plain text.
"""

import io
import logging
import re

from minio import Minio

from app.core.config import settings
from app.integrations.resume_parser.protocol import ResumeSource
from app.services import storage_service

logger = logging.getLogger(__name__)

#: Below this, treat the document as having no usable text layer.
MIN_USABLE_CHARS = 120


class ResumeTextError(Exception):
    """Raised when no usable text can be pulled out of the upload."""


def _collapse_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages[:20]:  # a CV past 20 pages is not a CV
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001 — one bad page shouldn't lose the rest
            logger.warning("pdf page extraction failed: %s", exc)
    return "\n\n".join(pages)


def _from_docx(data: bytes) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _client() -> Minio:
    return storage_service.internal_client()


def download(object_key: str) -> bytes:
    response = _client().get_object(settings.storage_bucket, object_key)
    try:
        return response.read()
    finally:
        response.close()
        response.release_conn()


def build_source(object_key: str, content_type: str) -> ResumeSource:
    """Download the upload and package it for the parser.

    A PDF is handed over as bytes. Extracting its text here would be actively
    harmful: a CV is a laid-out document, and flattening two columns into one
    stream is what makes a sidebar's skill list look like it belongs to the job
    printed next to it. The model reads the page instead.

    Text is still attached when a text layer exists, because the stub parser
    cannot see and `raw_text` is populated from it — but for a PDF its absence
    is no longer fatal. Scanned CVs, which used to fail outright, now parse.
    """
    data = download(object_key)
    is_pdf = content_type == "application/pdf" or object_key.lower().endswith(".pdf")

    if is_pdf:
        try:
            text = _collapse_whitespace(_from_pdf(data))
        except Exception as exc:  # noqa: BLE001 — the PDF itself is what matters
            logger.warning("pdf text layer unreadable, sending document only: %s", exc)
            text = ""
        return ResumeSource(
            text=text or None,
            pdf=data,
            filename=object_key.rsplit("/", 1)[-1] or "resume.pdf",
        )

    if "wordprocessingml" in content_type or object_key.lower().endswith(".docx"):
        text = _from_docx(data)
    elif content_type.startswith("text/"):
        text = data.decode("utf-8", errors="replace")
    elif content_type.startswith("image/"):
        raise ResumeTextError(
            "Image files can't be read automatically. Upload a PDF or Word "
            "document, or fill in the form manually."
        )
    else:
        raise ResumeTextError(f"Unsupported file type for parsing: {content_type}")

    text = _collapse_whitespace(text)

    # Only non-PDF formats still depend on extraction, so this is the one path
    # where empty text genuinely means there is nothing to parse.
    if len(text) < MIN_USABLE_CHARS:
        raise ResumeTextError(
            "No readable text found in this document. Upload a PDF, or fill in "
            "the form manually."
        )

    return ResumeSource(text=text)
